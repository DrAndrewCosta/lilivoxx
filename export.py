
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

def exportar_laudo_docx(texto, nome_paciente, tipo_exame, data_exame_iso, medico, template_path, nome_arquivo="laudo_final.docx", data_nascimento=None):
    doc = Document()

    # Título e frase técnica do template
    template_doc = Document(template_path)
    parags = [p.text.strip() for p in template_doc.paragraphs if p.text.strip()]
    titulo_template = parags[0] if parags else tipo_exame
    frase_tecnica = parags[1] if len(parags) > 1 else ""

    doc.add_heading(titulo_template, level=1)
    if frase_tecnica:
        doc.add_paragraph(frase_tecnica)

    doc.add_paragraph(f"Paciente: {nome_paciente}")
    doc.add_paragraph(f"Data do exame: {data_exame_iso}")
    if data_nascimento:
        doc.add_paragraph(f"Data de nascimento: {data_nascimento}")
    if medico:
        doc.add_paragraph(f"Médico solicitante: {medico}")
    doc.add_paragraph("")

    for linha in texto.split("\n"):
        paragrafo = doc.add_paragraph()
        run = paragrafo.add_run(linha)
        run.font.size = Pt(11)

    doc.save(nome_arquivo)
    return nome_arquivo

def exportar_laudo_pdf(nome_paciente, data_exame, texto, medico=None, data_nascimento=None):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    pdf.set_font("Arial", style="B", size=14)
    pdf.cell(200, 10, txt="Laudo Médico", ln=True, align="C")

    pdf.set_font("Arial", size=12)
    pdf.ln(10)
    pdf.cell(200, 10, txt=f"Paciente: {nome_paciente}", ln=True)
    pdf.cell(200, 10, txt=f"Data do exame: {data_exame}", ln=True)
    if data_nascimento:
        pdf.cell(200, 10, txt=f"Data de nascimento: {data_nascimento}", ln=True)
    if medico:
        pdf.cell(200, 10, txt=f"Médico solicitante: {medico}", ln=True)
    pdf.ln(5)

    for linha in texto.split("\n"):
        pdf.multi_cell(0, 8, txt=linha)

    nome_arquivo = f"laudo_final_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    pdf.output(nome_arquivo)
    return nome_arquivo
